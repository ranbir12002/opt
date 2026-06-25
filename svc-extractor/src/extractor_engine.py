"""
extractor_engine.py

Core extraction logic (NO FastAPI here).
- Routes by file type
- Extracts text/tables/key-values (lightweight)
- Computes: is_useful, confidence, warnings
- Includes PDF (native + OCR fallback) because your data is mostly PDF

Dependencies (pip):
  pymupdf
  pdfplumber
  pillow
  pytesseract
  pandas
  openpyxl
  python-docx

Optional (recommended for better OCR preprocessing):
  opencv-python

Runtime dependency (if you enable OCR):
  Tesseract installed on OS, and pytesseract can find it.

Design goals:
- Deterministic output
- Never crash the service: return status="partial"/"failed" with warnings/errors
- "Usefulness" + confidence computed from measurable signals
"""

from __future__ import annotations

import io
import re
import math
import hashlib
from dataclasses import dataclass, asdict
from typing import Any, Dict, List, Optional, Tuple, Literal

# -----------------------------
# Config / thresholds (tunable)
# -----------------------------

DOC_TYPE_HINTS = {"unknown", "invoice", "sop", "work_order", "generic"}

MIN_TEXT_CHARS_BY_TYPE = {
    "csv": 50,
    "xlsx": 50,
    "docx": 100,
    "pdf": 200,       # total text chars (or per-page rule below)
    "image": 80,
}

MIN_TABLE_ROWS_BY_TYPE = {
    "csv": 5,
    "xlsx": 5,
    "docx": 1,
    "pdf": 1,         # table extraction is best-effort
    "image": 0,
}

# PDF scanned detection
PDF_SCANNED_AVG_CHARS_PER_PAGE_THRESHOLD = 25  # < 25 chars/page → likely scanned or blocked
PDF_WEAK_AVG_CHARS_PER_PAGE_THRESHOLD = 80     # 25–80 → weak extraction
PDF_MAX_OCR_PAGES_DEFAULT = 10                 # safety limit for OCR runtime (tune)

# Confidence thresholds for downstream routing/fallback
CONFIDENCE_USEFUL_DEFAULT = 0.65

# If doc_type_hint is invoice/sop, require minimal key fields
MIN_KV_HITS_INVOICE = 2
MIN_KV_HITS_SOP = 1  # SOP is mainly text; KV is optional but we keep small requirement


# -----------------------------
# Data structures
# -----------------------------

Status = Literal["ok", "partial", "failed"]

@dataclass
class PageText:
    page: int
    text: str
    confidence: float = 0.0
    method: str = "native"  # native|ocr|mixed

@dataclass
class TableBlock:
    name: str
    page: Optional[int]
    rows: List[List[str]]
    confidence: float = 0.0
    method: str = "native"  # native|ocr

@dataclass
class KeyValueHit:
    key: str
    value: str
    confidence: float
    evidence: Optional[str] = None
    page: Optional[int] = None

@dataclass
class ExtractionResult:
    file_id: str
    filename: str
    detected_type: str
    doc_type_hint: str

    status: Status
    is_useful: bool
    confidence: float

    text: str
    pages: List[PageText]
    tables: List[TableBlock]
    key_values: Dict[str, KeyValueHit]

    warnings: List[str]
    errors: List[str]

    meta: Dict[str, Any]


# -----------------------------
# Public entrypoint
# -----------------------------

def run_extraction(
    file_bytes: bytes,
    filename: str,
    content_type: Optional[str] = None,
    doc_type_hint: str = "unknown",
    enable_ocr: bool = True,
    max_ocr_pages: int = PDF_MAX_OCR_PAGES_DEFAULT,
) -> ExtractionResult:
    """
    Main entrypoint used by extractor_service.py

    Returns a normalized ExtractionResult with:
      - status: ok|partial|failed
      - is_useful: True/False
      - confidence: 0–1
      - warnings/errors for downstream logic + fallback
    """
    hint = (doc_type_hint or "unknown").lower().strip()
    if hint not in DOC_TYPE_HINTS:
        hint = "unknown"

    file_id = _sha256_id(file_bytes)
    detected = detect_file_type(filename, content_type)

    warnings: List[str] = []
    errors: List[str] = []

    pages: List[PageText] = []
    tables: List[TableBlock] = []
    kv: Dict[str, KeyValueHit] = {}
    full_text = ""

    status: Status = "ok"

    try:
        if detected == "pdf":
            full_text, pages, tables, kv, warnings, errors, meta = _extract_pdf(
                file_bytes=file_bytes,
                filename=filename,
                doc_type_hint=hint,
                enable_ocr=enable_ocr,
                max_ocr_pages=max_ocr_pages,
            )
        elif detected == "xlsx":
            full_text, tables, kv, warnings, errors, meta = _extract_xlsx(file_bytes, filename, hint)
        elif detected == "csv":
            full_text, tables, kv, warnings, errors, meta = _extract_csv(file_bytes, filename, hint)
        elif detected == "docx":
            full_text, tables, kv, warnings, errors, meta = _extract_docx(file_bytes, filename, hint)
        elif detected == "image":
            full_text, pages, kv, warnings, errors, meta = _extract_image(
                file_bytes=file_bytes,
                filename=filename,
                doc_type_hint=hint,
                enable_ocr=enable_ocr,
            )
        else:
            status = "failed"
            errors.append(f"Unsupported file type: {detected}")
            meta = {}

    except Exception as e:
        status = "failed"
        errors.append(f"Unhandled extraction error: {type(e).__name__}: {e}")
        meta = {}

    # Normalize outputs
    full_text = normalize_text(full_text)
    for p in pages:
        p.text = normalize_text(p.text)

    # Compute metrics
    metrics = _compute_metrics(
        detected_type=detected,
        doc_type_hint=hint,
        text=full_text,
        tables=tables,
        key_values=kv,
        pages=pages,
        errors=errors,
        warnings=warnings,
    )

    # Decide is_useful + confidence
    is_useful = _decide_useful(metrics)
    confidence = _compute_confidence(metrics)

    # Add warnings from quality
    warnings.extend(metrics.get("quality_warnings", []))
    warnings = _dedupe_preserve_order(warnings)

    # Decide final status
    if status != "failed":
        if errors:
            status = "partial"
        else:
            status = "ok"

    return ExtractionResult(
        file_id=file_id,
        filename=filename,
        detected_type=detected,
        doc_type_hint=hint,
        status=status,
        is_useful=is_useful,
        confidence=confidence,
        text=full_text,
        pages=pages,
        tables=tables,
        key_values=kv,
        warnings=warnings,
        errors=errors,
        meta=meta,
    )


# -----------------------------
# File type detection
# -----------------------------

def detect_file_type(filename: str, content_type: Optional[str]) -> str:
    name = (filename or "").lower().strip()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    # MIME hints
    ct = (content_type or "").lower()

    if ext in {"pdf"} or "pdf" in ct:
        return "pdf"
    if ext in {"xlsx", "xls"} or "spreadsheet" in ct or "excel" in ct:
        return "xlsx"
    if ext in {"csv"} or "csv" in ct:
        return "csv"
    if ext in {"docx"} or "word" in ct:
        return "docx"
    if ext in {"png", "jpg", "jpeg", "webp", "bmp", "tiff"} or ct.startswith("image/"):
        return "image"

    # default unknown
    return "unknown"


# -----------------------------
# PDF extraction (native + OCR)
# -----------------------------

def _extract_pdf(
    file_bytes: bytes,
    filename: str,
    doc_type_hint: str,
    enable_ocr: bool,
    max_ocr_pages: int,
) -> Tuple[str, List[PageText], List[TableBlock], Dict[str, KeyValueHit], List[str], List[str], Dict[str, Any]]:
    warnings: List[str] = []
    errors: List[str] = []
    meta: Dict[str, Any] = {"method_attempts": []}

    pages: List[PageText] = []
    tables: List[TableBlock] = []
    full_text_parts: List[str] = []

    # 1) Native text extraction (PyMuPDF preferred)
    native_ok = False
    native_text_total = 0
    page_count = 0

    try:
        import fitz  # PyMuPDF
        meta["method_attempts"].append("pymupdf_text")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = doc.page_count

        for i in range(page_count):
            page = doc.load_page(i)
            text = page.get_text("text") or ""
            text = text.strip()
            native_text_total += len(_strip_ws(text))
            full_text_parts.append(text)
            pages.append(PageText(page=i + 1, text=text, confidence=0.0, method="native"))

        native_ok = True

    except Exception as e:
        errors.append(f"PDF native text (PyMuPDF) failed: {type(e).__name__}: {e}")
        native_ok = False

    # 2) If native failed completely, try pdfplumber
    if (not native_ok) or (native_text_total == 0 and page_count == 0):
        try:
            import pdfplumber
            meta["method_attempts"].append("pdfplumber_text")
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                page_count = len(pdf.pages)
                pages = []
                full_text_parts = []
                native_text_total = 0

                for i, p in enumerate(pdf.pages):
                    t = (p.extract_text() or "").strip()
                    native_text_total += len(_strip_ws(t))
                    full_text_parts.append(t)
                    pages.append(PageText(page=i + 1, text=t, confidence=0.0, method="native"))

                native_ok = True
        except Exception as e:
            errors.append(f"PDF native text (pdfplumber) failed: {type(e).__name__}: {e}")
            native_ok = False

    full_text = "\n\n".join([t for t in full_text_parts if t])

    # 3) Decide if scanned/weak
    avg_chars_per_page = _safe_div(native_text_total, max(page_count, 1))
    meta["pdf_pages"] = page_count
    meta["pdf_native_text_chars"] = native_text_total
    meta["pdf_avg_chars_per_page"] = avg_chars_per_page

    scanned_likely = avg_chars_per_page < PDF_SCANNED_AVG_CHARS_PER_PAGE_THRESHOLD
    weak_likely = PDF_SCANNED_AVG_CHARS_PER_PAGE_THRESHOLD <= avg_chars_per_page < PDF_WEAK_AVG_CHARS_PER_PAGE_THRESHOLD

    if scanned_likely:
        warnings.append("pdf_scanned_detected")
    elif weak_likely:
        warnings.append("pdf_weak_text_extraction")

    # 4) Table extraction (best-effort) using pdfplumber
    #    NOTE: not guaranteed. We keep it as extra signal; invoice agents can still parse line items from text.
    try:
        import pdfplumber
        meta["method_attempts"].append("pdfplumber_tables")
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, p in enumerate(pdf.pages):
                # extract_tables returns list of tables: List[List[List[str]]]
                tbs = p.extract_tables() or []
                for ti, tb in enumerate(tbs):
                    rows = [[(c or "").strip() for c in row] for row in tb if row]
                    # discard tiny tables
                    if len(rows) >= 2 and max(len(r) for r in rows) >= 2:
                        tables.append(TableBlock(
                            name=f"pdf_table_p{i+1}_{ti+1}",
                            page=i + 1,
                            rows=rows,
                            confidence=0.0,
                            method="native",
                        ))
    except Exception:
        # tables are optional; don't fail extraction
        warnings.append("pdf_table_extraction_failed")

    # 5) OCR fallback if scanned/weak and enabled
    ocr_pages: List[PageText] = []
    if enable_ocr and (scanned_likely or weak_likely):
        try:
            ocr_text, ocr_pages, ocr_warnings, ocr_meta = _ocr_pdf_pages(
                file_bytes=file_bytes,
                max_pages=max_ocr_pages,
            )
            meta.update(ocr_meta)
            warnings.extend(ocr_warnings)
            if ocr_text:
                warnings.append("ocr_used")
                # merge OCR text: prefer native if native is decent, else use OCR
                if scanned_likely or native_text_total == 0:
                    full_text = ocr_text
                    pages = ocr_pages
                else:
                    # weak-likely: append OCR to fill gaps
                    full_text = (full_text + "\n\n" + ocr_text).strip()
                    # keep existing pages; add OCR pages with method="ocr"
                    pages.extend(ocr_pages)

        except Exception as e:
            warnings.append("ocr_failed")
            errors.append(f"OCR failed: {type(e).__name__}: {e}")

    # 6) Key-value (light rules) from extracted text
    kv = _extract_kv_light(full_text, doc_type_hint)

    return full_text, pages, tables, kv, _dedupe_preserve_order(warnings), errors, meta


def _ocr_pdf_pages(file_bytes: bytes, max_pages: int) -> Tuple[str, List[PageText], List[str], Dict[str, Any]]:
    """
    Converts PDF pages to images and OCRs them.
    Uses PyMuPDF for rendering.
    """
    warnings: List[str] = []
    meta: Dict[str, Any] = {"ocr_max_pages": max_pages}
    pages: List[PageText] = []
    text_parts: List[str] = []

    try:
        import fitz  # PyMuPDF
    except Exception as e:
        raise RuntimeError("PyMuPDF required for PDF OCR rendering") from e

    # OCR libraries
    try:
        import pytesseract
        from PIL import Image
    except Exception as e:
        raise RuntimeError("pytesseract + Pillow required for OCR") from e

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = doc.page_count
    meta["pdf_total_pages_for_ocr"] = total_pages

    pages_to_ocr = min(total_pages, max_pages)
    if total_pages > max_pages:
        warnings.append("ocr_page_limit_applied")

    for i in range(pages_to_ocr):
        page = doc.load_page(i)
        # Render at higher DPI for better OCR
        zoom = 2.0  # 2.0 ~ 144 DPI if base 72
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes("png")

        img = Image.open(io.BytesIO(img_bytes))
        img = _preprocess_image_for_ocr(img)

        # OCR with confidence (TSV)
        ocr_text, avg_conf = _tesseract_ocr_with_confidence(img)
        ocr_text = (ocr_text or "").strip()

        if ocr_text:
            text_parts.append(ocr_text)

        pages.append(PageText(page=i + 1, text=ocr_text, confidence=avg_conf, method="ocr"))

    full_text = "\n\n".join([t for t in text_parts if t]).strip()
    meta["ocr_text_chars"] = len(_strip_ws(full_text))
    return full_text, pages, _dedupe_preserve_order(warnings), meta


# -----------------------------
# Image extraction (OCR)
# -----------------------------

def _extract_image(
    file_bytes: bytes,
    filename: str,
    doc_type_hint: str,
    enable_ocr: bool,
) -> Tuple[str, List[PageText], Dict[str, KeyValueHit], List[str], List[str], Dict[str, Any]]:
    warnings: List[str] = []
    errors: List[str] = []
    meta: Dict[str, Any] = {}

    if not enable_ocr:
        warnings.append("ocr_disabled")
        return "", [], {}, warnings, errors, meta

    try:
        from PIL import Image
    except Exception as e:
        errors.append(f"Pillow missing: {e}")
        return "", [], {}, warnings, errors, meta

    img = Image.open(io.BytesIO(file_bytes))
    meta["image_size"] = {"width": img.size[0], "height": img.size[1]}
    if img.size[0] < 800:
        warnings.append("image_low_resolution")

    img = _preprocess_image_for_ocr(img)
    text, avg_conf = _tesseract_ocr_with_confidence(img)
    text = (text or "").strip()

    pages = [PageText(page=1, text=text, confidence=avg_conf, method="ocr")]
    kv = _extract_kv_light(text, doc_type_hint)

    if avg_conf < 0.55:
        warnings.append("ocr_low_confidence")

    return text, pages, kv, _dedupe_preserve_order(warnings), errors, meta


def _preprocess_image_for_ocr(img):
    """
    Light preprocessing.
    If opencv is available, do better threshold + deskew.
    """
    from PIL import ImageOps

    # Convert to grayscale
    img = ImageOps.grayscale(img)

    # Try OpenCV-based preprocessing if available
    try:
        import cv2
        import numpy as np

        arr = np.array(img)

        # Denoise lightly
        arr = cv2.fastNlMeansDenoising(arr, None, h=10, templateWindowSize=7, searchWindowSize=21)

        # Adaptive threshold
        arr = cv2.adaptiveThreshold(arr, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                    cv2.THRESH_BINARY, 31, 10)

        # Deskew (simple)
        coords = np.column_stack(np.where(arr < 255))
        if coords.size > 0:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            (h, w) = arr.shape[:2]
            M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
            arr = cv2.warpAffine(arr, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

        from PIL import Image
        return Image.fromarray(arr)

    except Exception:
        # Fallback: simple autocontrast
        img = ImageOps.autocontrast(img)
        return img


def _tesseract_ocr_with_confidence(img) -> Tuple[str, float]:
    """
    Returns text + average confidence (0..1).
    """
    import pytesseract

    # OCR text
    text = pytesseract.image_to_string(img) or ""

    # OCR confidence using TSV
    avg_conf = 0.0
    try:
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        confs = []
        for c in data.get("conf", []):
            try:
                ci = float(c)
                if ci >= 0:
                    confs.append(ci)
            except Exception:
                pass
        if confs:
            avg_conf = sum(confs) / (len(confs) * 100.0)
        else:
            avg_conf = 0.5  # unknown but not zero
    except Exception:
        avg_conf = 0.5

    return text, float(max(0.0, min(1.0, avg_conf)))


# -----------------------------
# CSV / XLSX / DOCX
# -----------------------------

def _extract_csv(
    file_bytes: bytes,
    filename: str,
    doc_type_hint: str,
) -> Tuple[str, List[TableBlock], Dict[str, KeyValueHit], List[str], List[str], Dict[str, Any]]:
    warnings: List[str] = []
    errors: List[str] = []
    meta: Dict[str, Any] = {}

    # Try pandas first for robustness
    try:
        import pandas as pd
        bio = io.BytesIO(file_bytes)

        # encoding fallback strategy
        for enc in ("utf-8", "utf-8-sig", "cp1252", "latin1"):
            bio.seek(0)
            try:
                df = pd.read_csv(bio, encoding=enc)
                if enc != "utf-8":
                    warnings.append("encoding_guessed")
                break
            except Exception:
                df = None
        if df is None:
            raise ValueError("CSV decoding failed with common encodings")

        meta["csv_rows"] = int(df.shape[0])
        meta["csv_cols"] = int(df.shape[1])
        if df.shape[0] > 5000:
            warnings.append("large_file_detected")

        rows = [df.columns.tolist()] + df.fillna("").astype(str).values.tolist()
        tables = [TableBlock(name="csv_table_1", page=None, rows=rows, confidence=0.0, method="native")]

        # Also create a text view (helpful for LLM or debugging)
        text = df.head(200).to_csv(index=False)

        kv = _extract_kv_light(text, doc_type_hint)
        return text, tables, kv, _dedupe_preserve_order(warnings), errors, meta

    except Exception as e:
        errors.append(f"CSV parse failed: {type(e).__name__}: {e}")
        return "", [], {}, _dedupe_preserve_order(warnings), errors, meta


def _extract_xlsx(
    file_bytes: bytes,
    filename: str,
    doc_type_hint: str,
) -> Tuple[str, List[TableBlock], Dict[str, KeyValueHit], List[str], List[str], Dict[str, Any]]:
    warnings: List[str] = []
    errors: List[str] = []
    meta: Dict[str, Any] = {}

    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        sheets = wb.sheetnames
        meta["xlsx_sheets"] = sheets

        all_tables: List[TableBlock] = []
        text_parts: List[str] = []

        for sname in sheets:
            ws = wb[sname]
            # Read a bounded region to keep service fast
            max_rows = ws.max_row or 0
            max_cols = min(ws.max_column or 0, 50)

            rows: List[List[str]] = []
            merged = bool(ws.merged_cells.ranges)
            if merged:
                warnings.append("merged_cells_detected")

            for r in range(1, max_rows + 1):
                row_vals = []
                empty_row = True
                for c in range(1, max_cols + 1):
                    v = ws.cell(row=r, column=c).value
                    sv = "" if v is None else str(v)
                    if sv.strip():
                        empty_row = False
                    row_vals.append(sv)
                # keep a little slack: skip fully empty rows deep in sheet
                if empty_row and r > 50:
                    continue
                rows.append(row_vals)

            # Only keep non-trivial sheets
            non_empty_rows = sum(1 for r in rows if any((x or "").strip() for x in r))
            if non_empty_rows > 5000:
                warnings.append("large_sheet_detected")
            if non_empty_rows >= 2:
                all_tables.append(TableBlock(name=f"xlsx_{sname}", page=None, rows=rows, confidence=0.0, method="native"))
                # add a small text summary
                text_parts.append(f"--- SHEET: {sname} ---\n" + "\n".join([", ".join(r) for r in rows[:50]]))

        text = "\n\n".join(text_parts)
        kv = _extract_kv_light(text, doc_type_hint)
        return text, all_tables, kv, _dedupe_preserve_order(warnings), errors, meta

    except Exception as e:
        errors.append(f"XLSX parse failed: {type(e).__name__}: {e}")
        return "", [], {}, _dedupe_preserve_order(warnings), errors, meta


def _extract_docx(
    file_bytes: bytes,
    filename: str,
    doc_type_hint: str,
) -> Tuple[str, List[TableBlock], Dict[str, KeyValueHit], List[str], List[str], Dict[str, Any]]:
    warnings: List[str] = []
    errors: List[str] = []
    meta: Dict[str, Any] = {}

    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))

        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paras)
        meta["docx_paragraphs"] = len(paras)

        tables: List[TableBlock] = []
        for ti, t in enumerate(doc.tables):
            rows: List[List[str]] = []
            for r in t.rows:
                rows.append([cell.text.strip() for cell in r.cells])
            if len(rows) >= 1:
                tables.append(TableBlock(name=f"docx_table_{ti+1}", page=None, rows=rows, confidence=0.0, method="native"))

        if not text and tables:
            warnings.append("docx_no_paragraph_text_tables_only")
        elif not text and not tables:
            warnings.append("docx_empty_or_images_only")

        kv = _extract_kv_light(text, doc_type_hint)
        return text, tables, kv, _dedupe_preserve_order(warnings), errors, meta

    except Exception as e:
        errors.append(f"DOCX parse failed: {type(e).__name__}: {e}")
        return "", [], {}, _dedupe_preserve_order(warnings), errors, meta


# -----------------------------
# Light key-value extraction (optional)
# -----------------------------

def _extract_kv_light(text: str, doc_type_hint: str) -> Dict[str, KeyValueHit]:
    """
    Lightweight, deterministic key-field hints.
    Not meant to be perfect. Helps invoice agents and usefulness scoring.
    """
    t = text or ""
    kv: Dict[str, KeyValueHit] = {}

    # Work order / PO number patterns (tune to your conventions)
    wo_patterns = [
        r"\bWO[:\s\-#]*([A-Za-z0-9\-\/]+)\b",
        r"\bWork\s*Order[:\s\-#]*([A-Za-z0-9\-\/]+)\b",
        r"\bPO[:\s\-#]*([A-Za-z0-9\-\/]+)\b",
        r"\bPurchase\s*Order[:\s\-#]*([A-Za-z0-9\-\/]+)\b",
    ]
    wo = _first_group_match(t, wo_patterns)
    if wo:
        kv["work_order_or_po"] = KeyValueHit(key="work_order_or_po", value=wo, confidence=0.75, evidence="pattern_match")

    # Invoice number
    inv = _first_group_match(t, [r"\bInvoice\s*(No|Number)[:\s\-#]*([A-Za-z0-9\-\/]+)\b"])
    if inv:
        # inv returns group 2 if pattern has 2 groups; normalize:
        inv_val = inv if isinstance(inv, str) else str(inv)
        # If it returned "No" accidentally, try alternate:
        if inv_val.lower() in {"no", "number"}:
            inv_val = _first_group_match(t, [r"\bInvoice[:\s\-#]*([A-Za-z0-9\-\/]+)\b"]) or inv_val
        kv["invoice_no"] = KeyValueHit(key="invoice_no", value=inv_val, confidence=0.7, evidence="pattern_match")

    # Date (YYYY-MM-DD) common on generated docs; for AU invoices often DD/MM/YYYY too (we keep both)
    date_val = _first_group_match(t, [
        r"\b(20\d{2}[-\/]\d{2}[-\/]\d{2})\b",
        r"\b(\d{1,2}\/\d{1,2}\/20\d{2})\b",
    ])
    if date_val:
        kv["date_any"] = KeyValueHit(key="date_any", value=date_val, confidence=0.6, evidence="pattern_match")

    # Total amount (very rough; your invoice agent can do proper parsing)
    total_val = _first_group_match(t, [
        r"\bTotal\s*(Inc|Including)?\s*(GST|Tax)?[:\s]*\$?\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]{2})?)\b",
        r"\bAmount\s*Due[:\s]*\$?\s*([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]{2})?)\b",
    ])
    if total_val:
        # total_val might capture group 3; we pick last numeric group
        num = _last_number_in_string(total_val)
        if num:
            kv["total_amount"] = KeyValueHit(key="total_amount", value=num, confidence=0.6, evidence="pattern_match")

    return kv


# -----------------------------
# Normalization
# -----------------------------

def normalize_text(text: str) -> str:
    if not text:
        return ""
    # Normalize line endings and whitespace
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim trailing spaces on lines
    text = "\n".join([ln.rstrip() for ln in text.split("\n")])
    return text.strip()


# -----------------------------
# Metrics -> usefulness + confidence
# -----------------------------

def _compute_metrics(
    detected_type: str,
    doc_type_hint: str,
    text: str,
    tables: List[TableBlock],
    key_values: Dict[str, KeyValueHit],
    pages: List[PageText],
    errors: List[str],
    warnings: List[str],
) -> Dict[str, Any]:
    txt = text or ""
    text_chars = len(_strip_ws(txt))
    table_rows = sum(max(0, len(tb.rows) - 1) for tb in tables) if tables else 0  # minus header row heuristic

    kv_hits = len(key_values) if key_values else 0

    # Noise heuristic: ratio of alnum to total
    total = max(1, len(txt))
    alnum = sum(1 for ch in txt if ch.isalnum())
    alnum_ratio = alnum / total

    # PDF page signal
    avg_chars_per_page = None
    if detected_type == "pdf" and pages:
        per_page = [len(_strip_ws(p.text)) for p in pages[: min(len(pages), 50)]]
        avg_chars_per_page = sum(per_page) / max(1, len(per_page))

    quality_warnings: List[str] = []
    if text_chars < max(10, MIN_TEXT_CHARS_BY_TYPE.get(detected_type, 50) // 2) and table_rows == 0:
        quality_warnings.append("low_text_extracted")

    if alnum_ratio < 0.35 and text_chars > 0:
        quality_warnings.append("text_looks_noisy")

    if detected_type == "pdf" and avg_chars_per_page is not None:
        if avg_chars_per_page < PDF_SCANNED_AVG_CHARS_PER_PAGE_THRESHOLD:
            # already covered by pdf extractor warning, but keep consistent
            pass
        elif avg_chars_per_page < PDF_WEAK_AVG_CHARS_PER_PAGE_THRESHOLD:
            quality_warnings.append("pdf_layout_may_be_complex")

    return {
        "detected_type": detected_type,
        "doc_type_hint": doc_type_hint,
        "text_chars": text_chars,
        "table_rows": table_rows,
        "kv_hits": kv_hits,
        "alnum_ratio": alnum_ratio,
        "avg_chars_per_page": avg_chars_per_page,
        "errors_count": len(errors),
        "warnings_count": len(warnings),
        "quality_warnings": quality_warnings,
        "ocr_used": ("ocr_used" in warnings),
        "ocr_low_conf": ("ocr_low_confidence" in warnings),
        "pages_count": len(pages) if pages else 0,
    }


def _decide_useful(metrics: Dict[str, Any]) -> bool:
    detected = metrics["detected_type"]
    hint = metrics["doc_type_hint"]
    text_chars = metrics["text_chars"]
    table_rows = metrics["table_rows"]
    kv_hits = metrics["kv_hits"]
    errors_count = metrics["errors_count"]

    if detected == "unknown":
        return False
    if errors_count > 0 and text_chars == 0 and table_rows == 0:
        return False

    # Generic gates
    min_text = MIN_TEXT_CHARS_BY_TYPE.get(detected, 50)
    min_rows = MIN_TABLE_ROWS_BY_TYPE.get(detected, 0)

    passes_generic = (text_chars >= min_text) or (table_rows >= min_rows) or (kv_hits >= 1)

    if not passes_generic:
        return False

    # Hint-specific gates
    if hint == "invoice":
        return kv_hits >= MIN_KV_HITS_INVOICE or (text_chars >= (min_text * 2) and table_rows >= 1)
    if hint == "sop":
        return text_chars >= 300 or kv_hits >= MIN_KV_HITS_SOP

    return True


def _compute_confidence(metrics: Dict[str, Any]) -> float:
    detected = metrics["detected_type"]
    hint = metrics["doc_type_hint"]

    text_chars = metrics["text_chars"]
    table_rows = metrics["table_rows"]
    kv_hits = metrics["kv_hits"]
    alnum_ratio = metrics["alnum_ratio"]
    ocr_used = metrics["ocr_used"]
    ocr_low_conf = metrics["ocr_low_conf"]
    errors_count = metrics["errors_count"]
    avg_cpp = metrics["avg_chars_per_page"]

    # S_complete: how complete is the extraction for that type?
    if detected in {"csv", "xlsx"}:
        s_complete = 0.2
        if table_rows >= 5:
            s_complete = 0.9
        elif table_rows >= 1:
            s_complete = 0.7
        if text_chars >= 200:
            s_complete = max(s_complete, 0.75)
    elif detected == "docx":
        s_complete = 0.3
        if text_chars >= 500:
            s_complete = 0.9
        elif text_chars >= 150:
            s_complete = 0.7
        elif table_rows >= 1:
            s_complete = 0.6
    elif detected == "pdf":
        s_complete = 0.25
        # Use avg chars per page if available (more robust)
        if avg_cpp is not None:
            if avg_cpp >= 150:
                s_complete = 0.9
            elif avg_cpp >= 80:
                s_complete = 0.75
            elif avg_cpp >= 25:
                s_complete = 0.55
            else:
                s_complete = 0.35
        else:
            if text_chars >= 1200:
                s_complete = 0.9
            elif text_chars >= 400:
                s_complete = 0.7
            elif text_chars >= 200:
                s_complete = 0.55
            else:
                s_complete = 0.35
        if table_rows >= 1:
            s_complete = max(s_complete, 0.6)
    elif detected == "image":
        s_complete = 0.35
        if text_chars >= 400:
            s_complete = 0.8
        elif text_chars >= 120:
            s_complete = 0.65
        elif text_chars >= 60:
            s_complete = 0.5
    else:
        s_complete = 0.2

    # S_noise: less noise -> higher
    # alnum_ratio ~ 0.45-0.75 for normal text; <0.35 often noisy OCR fragments
    if text_chars == 0:
        s_noise = 0.0
    else:
        s_noise = max(0.0, min(1.0, (alnum_ratio - 0.25) / 0.55))  # scale 0.25..0.80

    # S_struct: tables help
    if table_rows >= 50:
        s_struct = 0.9
    elif table_rows >= 10:
        s_struct = 0.75
    elif table_rows >= 1:
        s_struct = 0.55
    else:
        s_struct = 0.2

    # S_kv: only meaningful for invoice/sop hints
    if hint in {"invoice", "sop"}:
        if hint == "invoice":
            s_kv = min(1.0, kv_hits / 4.0)  # up to 4 key fields in our light kv
        else:
            s_kv = 0.6 if kv_hits >= 1 else 0.2
    else:
        s_kv = 0.3 if kv_hits >= 1 else 0.2

    # Weighted combine
    conf = 0.45 * s_complete + 0.25 * s_noise + 0.20 * s_struct + 0.10 * s_kv

    # Penalize OCR usage slightly unless OCR confidence is good (we only have low_conf warning in v1)
    if ocr_used:
        conf -= 0.05
    if ocr_low_conf:
        conf -= 0.10

    # Penalize errors
    if errors_count >= 1:
        conf -= 0.10
    if errors_count >= 2:
        conf -= 0.10

    # Clamp
    conf = float(max(0.0, min(1.0, conf)))
    return conf


# -----------------------------
# Helpers
# -----------------------------

def _sha256_id(b: bytes) -> str:
    h = hashlib.sha256()
    h.update(b)
    return "sha256:" + h.hexdigest()

def _strip_ws(s: str) -> str:
    return re.sub(r"\s+", "", s or "")

def _safe_div(a: float, b: float) -> float:
    if b == 0:
        return 0.0
    return float(a) / float(b)

def _dedupe_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in items:
        if x and x not in seen:
            out.append(x)
            seen.add(x)
    return out

def _first_group_match(text: str, patterns: List[str]) -> Optional[str]:
    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            # Return last group if multiple
            if m.lastindex:
                return m.group(m.lastindex)
            return m.group(0)
    return None

def _last_number_in_string(s: str) -> Optional[str]:
    if not s:
        return None
    nums = re.findall(r"[0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]{2})?", s)
    if not nums:
        return None
    return nums[-1].replace(",", "")

